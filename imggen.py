from re import sub
from string import punctuation
from os import listdir
from PIL import Image, ImageFont, ImageDraw
from matplotlib.pyplot import imshow
import numpy as np
from docx import Document


def main():
    # Options
    fontWidthScale = 0.8
    fontHeightScale = 1.1
    lineSpacing = 79
    kerning = 0.77
    textStartW = 105
    textStartH = 50

    files = list(filter(lambda x: x[-5:] == '.docx', listdir()))

    # Img config
    img = Image.open('tp.jpg')
    originWidth = img.width
    originHeight = img.height
    img = img.resize((int(img.width * fontHeightScale), int(img.height * fontWidthScale)))

    # Load font
    fontB = ImageFont.truetype('Montserrat-Bold.ttf', size=74)
    fontL = ImageFont.truetype('Montserrat-ExtraLight.ttf', size=35)

    draw = ImageDraw.Draw(img)

    # Text
    line1 = drawText(draw, (textStartW, textStartH), "DOMINGO DO LIVRAMENTO", font=fontB, spacing=int(fontB.size * kerning))
    line2 = drawText(draw, (textStartW, line1[1] + lineSpacing), "MATEUS 17 : 24 - 27", font=fontL, spacing=int(fontL.size * kerning))

    # Prepare img
    img = img.resize((originWidth, originHeight))

    i = 0
    for file in files:
        doc = Document(file)

        # i = 0
        # for paragraph in doc.paragraphs:
            # i += 1
        #     print(i, paragraph.text)

        for paragraph in doc.paragraphs:
            text = paragraph.text


            if any(x in sub(r'[^\w\s]', '', text[:3]).split() for x in ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9']) and not any(x.islower() for x in text[:10]):
                i += 1
                print(text + '\n')

    print("found:", i)


def drawText(drawer, xy, text, font, spacing):
    drawer.text(xy, text, font=font, spacing=spacing)
    return drawer.textbbox(xy, text, font=font, spacing=spacing)


def showText(document):
     i = 0
     for paragraph in document.paragraphs:
         i += 1
         print(i, paragraph.text)

def show(imgv):
    imshow(np.array(imgv))


if __name__ == '__main__':
    main()
